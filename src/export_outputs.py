from __future__ import annotations

from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches

from src.train_models import predict_probabilities


def save_prediction_frame(
    metadata: pd.DataFrame,
    probabilities: np.ndarray,
    threshold: float,
    model_name: str,
    file_path: Path,
) -> pd.DataFrame:
    prediction_frame = metadata.copy()
    prediction_frame["model"] = model_name
    prediction_frame["predicted_probability"] = probabilities
    prediction_frame["predicted_class"] = (probabilities >= threshold).astype(int)
    prediction_frame.to_csv(file_path, index=False)
    return prediction_frame


def export_susceptibility_with_coordinates(
    model: Any,
    feature_frame: pd.DataFrame,
    metadata: pd.DataFrame,
    threshold: float,
    model_name: str,
    file_path: Path,
) -> pd.DataFrame:
    probabilities = predict_probabilities(model, feature_frame)
    export_frame = metadata.copy()
    export_frame["model"] = model_name
    export_frame["predicted_probability"] = probabilities
    export_frame["predicted_class"] = (probabilities >= threshold).astype(int)
    export_frame.to_csv(file_path, index=False)
    return export_frame


def create_conference_docx(
    summary_frame: pd.DataFrame,
    test_metrics_frame: pd.DataFrame,
    validation_metrics_frame: pd.DataFrame,
    docs_dir: Path,
    figures_dir: Path,
) -> Path:
    docs_dir.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("Conference Paper Draft", level=0)
    document.add_paragraph(
        "Leakage-aware classical machine learning for flood susceptibility mapping under a strict spatial holdout."
    )

    document.add_heading("Abstract", level=1)
    document.add_paragraph(
        "This study evaluates classical machine learning models for flood susceptibility mapping "
        "using a tabular geospatial dataset and a strict spatial holdout design. Mixed predictor "
        "types were handled through categorical encoding of land use/land cover, cyclic encoding "
        "of aspect, missing-indicator engineering, median imputation, and reproducible preprocessing. "
        "Eight classical models were compared: Decision Tree, Random Forest, Extra Trees, AdaBoost, "
        "Gradient Boosting, HistGradientBoosting, XGBoost, and LightGBM. Under spatial evaluation, "
        "Random Forest achieved the highest test ROC-AUC, while XGBoost achieved the strongest average "
        "precision and F1 score, supporting tree-based ensembles as the most effective family for this dataset."
    )

    document.add_heading("Dataset Summary", level=1)
    for _, row in summary_frame.iterrows():
        document.add_paragraph(f"{row['metric']}: {row['value']}", style="List Bullet")

    document.add_heading("Validation Results", level=1)
    document.add_paragraph(
        "Top validation performance was achieved by XGBoost, LightGBM, and Random Forest under the spatial split."
    )
    _add_table(document, validation_metrics_frame.head(5))

    document.add_heading("Test Results", level=1)
    document.add_paragraph(
        "Random Forest produced the best ROC-AUC, while XGBoost produced the highest average precision and F1."
    )
    _add_table(document, test_metrics_frame)

    document.add_heading("Discussion", level=1)
    document.add_paragraph(
        "The results suggest that leakage-aware spatial validation is substantially more difficult than random splitting, "
        "which explains the moderate but realistic performance range. Tree-based ensembles remained the strongest family, "
        "likely because the dataset contains skewed continuous variables, categorical land-cover information, and nonlinear "
        "interactions that are well handled by ensemble trees."
    )

    comparison_figure = figures_dir / "model_comparison.png"
    if comparison_figure.exists():
        document.add_heading("Key Figure", level=1)
        document.add_picture(str(comparison_figure), width=Inches(6.5))

    output_path = docs_dir / "conference_paper.docx"
    document.save(output_path)
    return output_path


def _add_table(document: Document, frame: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, column in enumerate(frame.columns):
        header_cells[index].text = str(column)
    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(frame.columns):
            value = row[column]
            if isinstance(value, float):
                cells[index].text = f"{value:.4f}"
            else:
                cells[index].text = str(value)
